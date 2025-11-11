# rag/parser.py — TMU Weekly (state-machine parser, finalized)
from __future__ import annotations
import re
import datetime as dt
from typing import List, Dict, Optional, Tuple
from docx import Document

# Regex
RE_DOW_HDR    = re.compile(r"\b(Thứ\s*[2-7]|Chủ\s*nhật|CN|cn|thu\s*[2-7])\b", re.I)
RE_DDMM       = re.compile(r"\b(\d{1,2})[\/\-](\d{1,2})\b")
RE_DDMMYY     = re.compile(r"\b(\d{1,2})[\/\-](\d{1,2})[\/\-](\d{2,4})\b")

# giờ: 08:30 | 8h30 | 8:00 | 8h
RE_HHMM       = re.compile(r"\b(\d{1,2})[:h](\d{2})\b", re.I)
RE_HH         = re.compile(r"\b(\d{1,2})\s*h\b", re.I)

# khoảng giờ: 08:00-11:30 | 8h-11h30 | 08:00 đến 11:30 | Từ 8h đến 11h30
RE_TIME_RANGE = re.compile(
    r"\b(?:từ\s*)?(\d{1,2})(?:[:h](\d{2}))?\s*(?:-|–|—|đến|tới|->)\s*(\d{1,2})(?:[:h](\d{2}))?",
    re.I
)

# location tag & bullets
RE_LOC_TAG    = re.compile(r"\b(địa\s*điểm\s*[:：]|tại)\b", re.I)
RE_BULLET     = re.compile(r"^[\*\-\u2022]+\s*")

# TP / Thành phần / Mời dự
RE_TP         = re.compile(r"^(TP|Thành\s*phần|Mời\s*dự)\s*[:：\-]\s*(.+)$", re.I)

DOW_VI = ["Thứ 2","Thứ 3","Thứ 4","Thứ 5","Thứ 6","Thứ 7","Chủ nhật"]

# Helpers
def _fmt_date(d: dt.date) -> str:
    return f"{d.day:02d}/{d.month:02d}/{d.year:04d}"

def _dow_vi(d: dt.date) -> str:
    # datetime.weekday(): 0=Mon..6=Sun
    return DOW_VI[d.weekday()] if d.weekday() < 6 else "Chủ nhật"

def _coerce_year(day: int, mon: int, default_year: int) -> Optional[dt.date]:
    try:
        return dt.date(default_year, mon, day)
    except ValueError:
        return None

def _smart_cap(s: str) -> str:
    s = s.strip()
    if not s:
        return s
    return s[0].upper() + s[1:]

def _norm_time(s: str) -> Tuple[Optional[str], Optional[str]]:
    # Ưu tiên khoảng giờ
    m = RE_TIME_RANGE.search(s)
    if m:
        h1, m1, h2, m2 = m.groups()
        start = f"{int(h1):02d}:{int(m1 or 0):02d}"
        end   = f"{int(h2):02d}:{int(m2 or 0):02d}"
        return start, end

    # Một mốc giờ
    mm = RE_HHMM.search(s)
    if mm:
        h, m_ = mm.groups()
        return f"{int(h):02d}:{int(m_):02d}", None
    hh = RE_HH.search(s)
    if hh:
        return f"{int(hh.group(1)):02d}:00", None
    return None, None

def infer_year_from_doc(doc: Document) -> Optional[int]:
    # đoán năm xuất hiện trong file
    def _scan(text: str) -> Optional[int]:
        for m in re.finditer(r"\b(20\d{2})\b", text):
            y = int(m.group(1))
            if 2000 <= y <= 2100:
                return y
        return None
    y = _scan(" ".join(p.text for p in doc.paragraphs))
    if y: return y
    for tb in doc.tables:
        for r in tb.rows:
            y = _scan(" | ".join(c.text for c in r.cells))
            if y: return y
    return None

# Core Parser
def parse_docx_as_table(path: str, default_year: Optional[int] = None) -> List[Dict]:
    """
    Parser chuyên TMU:
      - Cột trái: 'Thứ X' + 'dd/mm'
      - Cột phải: từng dòng/bullet là một sự kiện
      - Dòng 'TP/Thành phần/Mời dự' ghép vào event trước đó
    Trả về: [{date,dow,start,end,location,participants,title,raw}]
    """
    doc  = Document(path)
    year = default_year or infer_year_from_doc(doc) or dt.date.today().year

    events: List[Dict] = []
    cur_date: Optional[str] = None
    cur_dow:  Optional[str] = None
    last_event_idx: Optional[int] = None

    def _scan_day_and_date(s: str) -> bool:
        nonlocal cur_date, cur_dow
        s1 = " ".join(s.split())
        m_dow = RE_DOW_HDR.search(s1)
        m_dm  = RE_DDMM.search(s1) or RE_DDMMYY.search(s1)
        if m_dow and m_dm:
            d = int(m_dm.group(1)); m = int(m_dm.group(2))
            y = int(m_dm.group(3)) + 2000 if (m_dm.lastindex == 3 and len(m_dm.group(3)) == 2) else year
            d_real = _coerce_year(d, m, y)
            if d_real:
                cur_date = _fmt_date(d_real)
                w = m_dow.group(1).lower()
                if "chủ nhật" in w or w in ("cn","CN"):
                    cur_dow = "Chủ nhật"
                else:
                    num = re.search(r"[2-7]", w)
                    cur_dow = f"Thứ {num.group(0)}" if num else None
                return True
        return False

    def _flush_tp(tp_text: str):
        nonlocal last_event_idx
        if last_event_idx is None:
            return
        tp_text = _smart_cap(tp_text.strip(" .;"))
        prev = events[last_event_idx].get("participants")
        events[last_event_idx]["participants"] = f"{prev}; {tp_text}" if prev else tp_text

    def _extract_location(full_text: str) -> Tuple[Optional[str], Optional[Tuple[int,int]]]:
        m = RE_LOC_TAG.search(full_text)
        if not m:
            return None, None
        tail = full_text[m.end():].strip()
        # cắt trước phần TP/Thành phần nếu có
        cut_at = None
        tp_pos = re.search(r"\b(Thành\s*phần|TP|Mời\s*dự)\b", tail, re.I)
        if tp_pos:
            cut_at = tp_pos.start()
        else:
            colon = tail.find(":")
            if 0 <= colon <= 40:
                cut_at = colon
        loc = (tail[:cut_at] if cut_at is not None else tail).strip(" .;–—|-")
        loc = _smart_cap(loc)
        return (loc or None), (m.start(), m.end())

    def _emit_event(raw_line: str):
        nonlocal last_event_idx
        raw = " ".join(raw_line.split())
        raw = RE_BULLET.sub("", raw)
        if not raw:
            return

        start, end = _norm_time(raw)
        location, loc_span = _extract_location(raw)

        # title: bỏ giờ + bỏ cụm "tại/địa điểm"
        title = RE_TIME_RANGE.sub("", raw)
        title = RE_HHMM.sub("", title)
        title = RE_HH.sub("", title)
        if loc_span:
            title = title[:loc_span[0]] + title[loc_span[1]:]
        title = _smart_cap(title.strip(" ,;–—|-"))

        ev = {
            "date": cur_date,
            "dow":  cur_dow,
            "start": start,
            "end": end,
            "location": location,
            "participants": None,
            "title": title if title else None,
            "raw": raw_line.strip()
        }
        events.append(ev)
        last_event_idx = len(events) - 1

    # Scan tài liệu
    if doc.tables:
        for tb in doc.tables:
            for row in tb.rows:
                left  = row.cells[0].text if len(row.cells) >= 1 else ""
                right = row.cells[1].text if len(row.cells) >= 2 else ""

                if left.strip():
                    _scan_day_and_date(left)

                for line in (l.strip() for l in right.split("\n")):
                    if not line:
                        continue
                    m_tp = RE_TP.match(line)
                    if m_tp:
                        _flush_tp(m_tp.group(2))
                        continue
                    if _scan_day_and_date(line):
                        continue
                    _emit_event(line)


    else:
        for p in doc.paragraphs:
            line = p.text.strip()
            if not line:
                continue
            m_tp = RE_TP.match(line)
            if m_tp:
                _flush_tp(m_tp.group(2))
                continue
            if _scan_day_and_date(line):
                continue
            _emit_event(line)

    # Lọc rác & bổ sung DOW từ date nếu thiếu
    events = [e for e in events if e.get("title") or e.get("start") or e.get("date")]
    for e in events:
        if not e.get("dow") and e.get("date"):
            try:
                d = dt.datetime.strptime(e["date"], "%d/%m/%Y").date()
                e["dow"] = _dow_vi(d)
            except Exception:
                pass

    return events