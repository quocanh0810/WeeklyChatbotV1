# rag/parser.py
from __future__ import annotations
import re, datetime as dt
from collections import Counter
from typing import List, Dict, Optional
from docx import Document

RE_DATE_YYYY = re.compile(r"\b(\d{1,2})[\/\-](\d{1,2})[\/\-](20\d{2})\b")
RE_DATE_DDMM = re.compile(r"\b(\d{1,2})[\/\-](\d{1,2})\b")
RE_TIME      = re.compile(r"\b(\d{1,2})(?::|h)(\d{2})\b")      # 08:30 | 8h30
RE_TIME_H    = re.compile(r"\b(\d{1,2})h\b")                   # 8h
RE_DOW       = re.compile(r"\b(thứ\s*[2-7]|thứ\s*cn|cn|chủ nhật|thu\s*[2-7])\b", re.I)

def _norm_dow(s: str) -> Optional[str]:
    s = s.strip().lower()
    s = s.replace("thu", "thứ").replace("  ", " ")
    m = re.search(RE_DOW, s)
    if not m: 
        return None
    t = m.group(1).lower().replace("  ", " ")
    t = t.replace("chủ nhật", "cn").replace("thứ cn", "cn")
    t = t.replace("thứ ", "Thứ ")
    if t == "cn": return "Chủ nhật"
    if t.startswith("thứ"):
        n = re.findall(r"\d", t)
        if n: return f"Thứ {n[0]}"
    return None

def infer_year_from_doc(doc: Document) -> Optional[int]:
    # lấy năm xuất hiện nhiều nhất trong doc
    years = Counter(re.findall(r"\b(20\d{2})\b", "\n".join(p.text for p in doc.paragraphs)))
    # quét cả bảng
    for tbl in doc.tables:
        for row in tbl.rows:
            cell_text = " | ".join(c.text for c in row.cells)
            years.update(re.findall(r"\b(20\d{2})\b", cell_text))
    if not years: 
        return None
    year, _ = years.most_common(1)[0]
    try:
        y = int(year)
        if 2000 <= y <= 2100:
            return y
    except Exception:
        pass
    return None

def _fmt_date(d: int, m: int, y: int) -> str:
    return f"{d:02d}/{m:02d}/{y:04d}"

def _extract_time(s: str) -> (Optional[str], Optional[str]):
    # lấy 1 khoảng giờ (đầu–cuối) nếu có
    times = RE_TIME.findall(s)
    if times:
        # nếu có >=2 mốc thì coi [0] là start, [1] là end
        if len(times) >= 2:
            h1, m1 = times[0]; h2, m2 = times[1]
            return f"{int(h1):02d}:{int(m1):02d}", f"{int(h2):02d}:{int(m2):02d}"
        else:
            h1, m1 = times[0]
            return f"{int(h1):02d}:{int(m1):02d}", None
    m = RE_TIME_H.search(s)
    if m:
        return f"{int(m.group(1)):02d}:00", None
    return None, None

def parse_docx_as_table(path: str, default_year: Optional[int] = None) -> List[Dict]:
    """
    Trích các hàng trong bảng thành events:
    {date, dow, start, end, location, participants, title, raw}
    """
    doc = Document(path)
    year = default_year or infer_year_from_doc(doc) or dt.date.today().year

    events: List[Dict] = []

    def emit_from_text(raw_line: str):
        nonlocal year
        raw = " ".join(raw_line.split())
        if not raw: 
            return
        # date
        dstr, dow = None, None
        m = RE_DATE_YYYY.search(raw)
        if m:
            d, mth, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
            dstr = _fmt_date(d, mth, y)
        else:
            m2 = RE_DATE_DDMM.search(raw)
            if m2:
                d, mth = int(m2.group(1)), int(m2.group(2))
                dstr = _fmt_date(d, mth, year)
        # dow
        dow = _norm_dow(raw) or None
        # time
        start, end = _extract_time(raw)
        # heuristics cho location / participants / title
        location = None
        participants = None
        title = None

        # tách theo các nhãn thường gặp
        low = raw.lower()
        # địa điểm
        for key in ["tại ", "địa điểm:", "địa điểm "]:
            if key in low:
                cut = low.index(key) + len(key)
                location = raw[cut:].split("  ")[0].strip()
                break
        # thành phần
        for key in ["thành phần:", "tp:", "thành phần "]:
            if key in low:
                cut = low.index(key) + len(key)
                participants = raw[cut:].strip()
                break
        # title
        title = raw

        events.append({
            "date": dstr,
            "dow": dow,
            "start": start,
            "end": end,
            "location": location,
            "participants": participants,
            "title": title,
            "raw": raw,
        })

    # Ưu tiên quét bảng
    if doc.tables:
        for tbl in doc.tables:
            for row in tbl.rows:
                # gộp cả hàng thành một dòng, nếu bảng đã tách cột thì càng chuẩn
                line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                emit_from_text(line)

    # fallback: không có bảng → quét paragraph
    else:
        for p in doc.paragraphs:
            emit_from_text(p.text)

    # lọc những dòng rỗng rãi quá
    events = [e for e in events if any([e.get("date"), e.get("start"), e.get("title")])]
    return events